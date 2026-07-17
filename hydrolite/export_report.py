from __future__ import annotations

from base64 import b64encode
from datetime import datetime, timezone
from html import escape
import json
from pathlib import Path
import zipfile
from typing import Any

import pandas as pd
import yaml

from hydrolite.__version__ import __app_name__, __version__


REPORT_FILENAMES = {
    "markdown": "project_report.md",
    "docx": "project_report.docx",
    "html": "project_report.html",
    "pdf": "project_report.pdf",
    "pdf_unavailable": "project_report_pdf_unavailable.md",
    "bundle": "project_report_bundle.zip",
}
REPORT_ASSET_NAMES = set(REPORT_FILENAMES.values())
FORBIDDEN_PARTS = {"external", ".git", ".venv", "env", "__pycache__", ".streamlit"}
FORBIDDEN_SUFFIXES = {".pt", ".pth", ".ckpt", ".onnx", ".key", ".pem"}
FORBIDDEN_NAME_FRAGMENTS = ("credential", "credentials", "service-account", "secret", "token")


def _project_path(project_dir: str | Path) -> Path:
    return Path(project_dir).expanduser().resolve()


def _read_yaml(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    except Exception as exc:
        return {"_read_error": str(exc)}
    return data if isinstance(data, dict) else {"_read_error": "YAML root is not a mapping."}


def _safe_read_excel(path: Path, sheet_name: str | int = 0, max_rows: int = 25) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_excel(path, sheet_name=sheet_name).head(max_rows)
    except Exception as exc:
        return pd.DataFrame({"error": [str(exc)], "file": [str(path)]})


def _safe_read_csv(path: Path, max_rows: int = 25) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    try:
        return pd.read_csv(path).head(max_rows)
    except Exception as exc:
        return pd.DataFrame({"error": [str(exc)], "file": [str(path)]})


def _read_text(path: Path, max_chars: int = 12000) -> str:
    if not path.exists():
        return ""
    try:
        text = path.read_text(encoding="utf-8")
    except Exception as exc:
        return f"Unable to read {path}: {exc}"
    return text[:max_chars]


def _hms_comparison_for_project(repo_root: Path, project: Path) -> Path | None:
    comparison = repo_root / "output" / "hec_hms_comparison"
    manifest = comparison / "comparison_manifest.json"
    if not manifest.is_file():
        return None
    try:
        payload = json.loads(manifest.read_text(encoding="utf-8"))
        source_project = Path(payload.get("hydrolite_project_dir", "")).expanduser().resolve()
    except (OSError, ValueError, json.JSONDecodeError):
        return None
    return comparison if source_project == project.resolve() else None


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value)


def _df_to_markdown(df: pd.DataFrame, empty_message: str = "unavailable") -> str:
    if df.empty:
        return empty_message
    try:
        return df.to_markdown(index=False)
    except Exception:
        return df.astype(str).to_csv(index=False)


def _df_to_html(df: pd.DataFrame) -> str:
    if df.empty:
        return "<p class=\"muted\">unavailable</p>"
    return df.to_html(index=False, classes="data-table", border=0, escape=True)


def _report_dir(project: Path) -> Path:
    reports = project / "reports"
    reports.mkdir(parents=True, exist_ok=True)
    return reports


def _relative_or_absolute(path: Path, base: Path) -> str:
    try:
        return str(path.relative_to(base))
    except ValueError:
        return str(path)


def _safe_bundle_member(path: Path, project: Path) -> bool:
    try:
        relative = path.resolve().relative_to(project)
    except ValueError:
        return False
    lowered = relative.as_posix().lower()
    parts = set(relative.parts)
    if parts & FORBIDDEN_PARTS:
        return False
    if path.suffix.lower() in FORBIDDEN_SUFFIXES:
        return False
    if any(fragment in lowered for fragment in FORBIDDEN_NAME_FRAGMENTS):
        return False
    if path.name == ".DS_Store" or path.suffix == ".pyc":
        return False
    return True


def _first_existing(paths: list[Path]) -> Path | None:
    for path in paths:
        if path.exists():
            return path
    return None


def _collect_case_swmm(project: Path) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    output_root = project / "output"
    if not output_root.exists():
        return pd.DataFrame()
    for case_dir in sorted(path for path in output_root.iterdir() if path.is_dir()):
        swmm = case_dir / "swmm" / "swmm_summary.xlsx"
        if not swmm.exists():
            continue
        row = _safe_read_excel(swmm).head(1)
        if row.empty:
            rows.append({"case_name": case_dir.name, "swmm_summary": str(swmm)})
        else:
            values = row.iloc[0].to_dict()
            values["case_name"] = case_dir.name
            values["swmm_summary"] = str(swmm)
            rows.append(values)
    return pd.DataFrame(rows)


def _collect_model_performance(project: Path) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    output_root = project / "output"
    if not output_root.exists():
        return pd.DataFrame()
    for case_dir in sorted(path for path in output_root.iterdir() if path.is_dir()):
        workbook = case_dir / "model_performance.xlsx"
        if not workbook.exists():
            continue
        metrics = _safe_read_excel(workbook, "metrics").head(1)
        row = {"case_name": case_dir.name, "model_performance_file": str(workbook)}
        if not metrics.empty:
            row.update(metrics.iloc[0].to_dict())
        rows.append(row)
    return pd.DataFrame(rows)


def list_report_assets(project_dir: str | Path) -> list[Path]:
    """Return generated report assets and source artifacts safe to expose/download."""
    project = _project_path(project_dir)
    candidates = [
        project / "project_summary.md",
        project / "reports" / "project_validation.xlsx",
        project / "reports" / "project_validation_report.md",
        project / "reports" / REPORT_FILENAMES["markdown"],
        project / "reports" / REPORT_FILENAMES["docx"],
        project / "reports" / REPORT_FILENAMES["html"],
        project / "reports" / REPORT_FILENAMES["pdf"],
        project / "reports" / REPORT_FILENAMES["pdf_unavailable"],
        project / "reports" / REPORT_FILENAMES["bundle"],
        project / "output" / "batch_summary.xlsx",
        project / "output" / "comparison" / "scenario_comparison.xlsx",
        project / "output" / "comparison" / "scenario_comparison.csv",
        project / "output" / "comparison" / "hydrolite_report.md",
        project / "output" / "comparison" / "peak_flow_comparison.png",
        project / "output" / "comparison" / "volume_comparison.png",
        project / "output" / "comparison" / "water_balance_comparison.png",
        project / "output" / "comparison" / "swmm_kpi_comparison.png",
    ]
    output = [path.resolve() for path in candidates if path.exists() and _safe_bundle_member(path, project)]
    return sorted(dict.fromkeys(output), key=lambda item: item.as_posix())


def collect_project_report_data(project_dir: str | Path) -> dict[str, Any]:
    project = _project_path(project_dir)
    repo_root = Path(__file__).resolve().parents[1]
    reports = _report_dir(project)
    config = _read_yaml(project / "project.yaml")
    comparison = project / "output" / "comparison" / "scenario_comparison.xlsx"
    hms_comparison_dir = _hms_comparison_for_project(repo_root, project) or project / "output" / "hec_hms_comparison"
    hms_comparison = hms_comparison_dir / "model_comparison_metrics.xlsx"
    validation = project / "reports" / "project_validation.xlsx"
    global_gee_summary = repo_root / "output" / "gee" / "gee_summary.xlsx"
    gee_summary = _first_existing([project / "output" / "gee" / "gee_summary.xlsx", global_gee_summary])
    openhydronet_report = _first_existing(
        [
            project / "output" / "openhydronet" / "inputs" / "openhydronet_input_report.md",
            Path("output/openhydronet/inputs/openhydronet_input_report.md").resolve(),
        ]
    )

    charts = [
        project / "output" / "comparison" / "peak_flow_comparison.png",
        project / "output" / "comparison" / "volume_comparison.png",
        project / "output" / "comparison" / "water_balance_comparison.png",
        project / "output" / "comparison" / "swmm_kpi_comparison.png",
        *sorted((hms_comparison_dir / "charts").glob("*.png")),
    ]
    existing_charts = [path for path in charts if path.exists()]
    expected = {
        "project_yaml": project / "project.yaml",
        "project_summary": project / "project_summary.md",
        "project_validation": validation,
        "scenario_comparison": comparison,
        "batch_summary": project / "output" / "batch_summary.xlsx",
        "gee_summary": gee_summary or project / "output" / "gee" / "gee_summary.xlsx",
        "openhydronet_report": openhydronet_report or project / "output" / "openhydronet" / "inputs" / "openhydronet_input_report.md",
    }
    missing = [
        {"asset": name, "path": str(path), "status": "missing"}
        for name, path in expected.items()
        if path is None or not Path(path).exists()
    ]

    data = {
        "project_dir": project,
        "reports_dir": reports,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "app_name": __app_name__,
        "version": __version__,
        "project": config,
        "project_name": config.get("project_name") or config.get("project_id") or project.name,
        "project_id": config.get("project_id") or project.name,
        "modules": config.get("modules") or {},
        "project_summary_text": _read_text(project / "project_summary.md"),
        "comparison_report_text": _read_text(project / "output" / "comparison" / "hydrolite_report.md"),
        "hms_comparison_report_text": _read_text(hms_comparison_dir / "comparison_report.md"),
        "openhydronet_report_text": _read_text(openhydronet_report) if openhydronet_report else "",
        "validation_overview": _safe_read_excel(validation, "case_overview"),
        "validation_errors": _safe_read_excel(validation, "case_errors"),
        "validation_warnings": _safe_read_excel(validation, "case_warnings"),
        "batch_summary": _safe_read_excel(project / "output" / "batch_summary.xlsx"),
        "overview": _safe_read_excel(comparison, "overview"),
        "hydrology_metrics": _safe_read_excel(comparison, "hydrology_metrics"),
        "water_balance_metrics": _safe_read_excel(comparison, "water_balance_metrics"),
        "swmm_metrics": _safe_read_excel(comparison, "swmm_metrics"),
        "coupling_metrics": _safe_read_excel(comparison, "coupling_metrics"),
        "performance_metrics": _safe_read_excel(comparison, "performance_metrics"),
        "missing_outputs": _safe_read_excel(comparison, "missing_outputs"),
        "gee_summary": _safe_read_excel(gee_summary) if gee_summary else pd.DataFrame(),
        "swmm_case_summaries": _collect_case_swmm(project),
        "model_performance": _collect_model_performance(project),
        "hms_comparison_summary": _safe_read_excel(hms_comparison, "summary"),
        "hms_comparison_metrics": _safe_read_excel(hms_comparison, "comparison_metrics"),
        "charts": existing_charts,
        "assets": list_report_assets(project),
        "missing_assets": pd.DataFrame(missing),
    }
    return data


def _executive_summary(data: dict[str, Any]) -> list[str]:
    overview = data["overview"]
    hydrology = data["hydrology_metrics"]
    water = data["water_balance_metrics"]
    swmm = data["swmm_metrics"]
    coupling = data["coupling_metrics"]
    lines = [f"Project `{data['project_name']}` report generated by {data['app_name']} {data['version']}."]
    if not overview.empty and "run_status" in overview.columns:
        success = int((overview["run_status"].astype(str).str.lower() == "success").sum())
        lines.append(f"Scenario overview includes {len(overview)} case(s), with {success} marked success.")
    if not hydrology.empty and "peak_flow" in hydrology.columns:
        values = pd.to_numeric(hydrology["peak_flow"], errors="coerce")
        if values.notna().any():
            idx = values.idxmax()
            lines.append(f"Highest peak flow case: `{hydrology.loc[idx, 'case_name']}` ({values.loc[idx]:.3f}).")
    if not water.empty:
        candidates: list[pd.Series] = []
        for column in ("max_subbasin_balance_error_percent", "outlet_balance_error_percent"):
            if column in water.columns:
                candidates.append(pd.to_numeric(water[column], errors="coerce").abs())
        if candidates:
            combined = pd.concat(candidates, axis=1).max(axis=1)
            if combined.notna().any():
                idx = combined.idxmax()
                lines.append(f"Largest water-balance error case: `{water.loc[idx, 'case_name']}` ({combined.loc[idx]:.3f}%).")
    if not swmm.empty and "swmm_status" in swmm.columns:
        statuses = sorted(set(swmm["swmm_status"].dropna().astype(str)))
        lines.append(f"SWMM status values found: {', '.join(statuses) if statuses else 'unavailable'}.")
    if not coupling.empty and "coupling_status" in coupling.columns:
        failed = coupling[coupling["coupling_status"].astype(str).str.lower().str.contains("fail", na=False)]
        if not failed.empty:
            lines.append("Coupling failures: " + ", ".join(failed["case_name"].astype(str).tolist()) + ".")
    if len(lines) == 1:
        lines.append("Several analysis outputs are missing; unavailable sections are documented in this report.")
    return lines


def render_project_report_markdown(project_dir: str | Path, output_path: str | Path | None = None) -> Path:
    data = collect_project_report_data(project_dir)
    output = Path(output_path).expanduser().resolve() if output_path else data["reports_dir"] / REPORT_FILENAMES["markdown"]
    output.parent.mkdir(parents=True, exist_ok=True)
    modules = data["modules"]
    lines = [
        f"# {data['project_name']} Project Report",
        "",
        f"- App: `{data['app_name']} {data['version']}`",
        f"- Generated at: `{data['generated_at']}`",
        f"- Project path: `{data['project_dir']}`",
        "",
        "## Executive Summary",
        "",
        *[f"- {line}" for line in _executive_summary(data)],
        "",
        "## Project Overview",
        "",
        f"- Project ID: `{data['project_id']}`",
        f"- Description: {data['project'].get('description', 'unavailable')}",
        "",
        "### Enabled Modules",
        "",
        *[f"- {name}: `{enabled}`" for name, enabled in modules.items()],
        "",
        "## Validation",
        "",
        _df_to_markdown(data["validation_overview"]),
        "",
        "### Validation Errors",
        "",
        _df_to_markdown(data["validation_errors"]),
        "",
        "### Validation Warnings",
        "",
        _df_to_markdown(data["validation_warnings"]),
        "",
        "## Scenario Runs",
        "",
        _df_to_markdown(data["batch_summary"]),
        "",
        "## Hydrology and Routing Comparison",
        "",
        _df_to_markdown(data["hydrology_metrics"]),
        "",
        "## Water Balance",
        "",
        _df_to_markdown(data["water_balance_metrics"]),
        "",
        "## SWMM Results",
        "",
        _df_to_markdown(data["swmm_metrics"]),
        "",
        "## HydroLite-SWMM Coupling",
        "",
        _df_to_markdown(data["coupling_metrics"]),
        "",
        "## GEE Data Center",
        "",
        _df_to_markdown(data["gee_summary"]),
        "",
        "## OpenHydroNet Input Summary",
        "",
        data["openhydronet_report_text"] or "unavailable",
        "",
        "## Observed Flow Evaluation",
        "",
        _df_to_markdown(data["model_performance"]),
        "",
        "## HEC-HMS and HydroLite Event Comparison",
        "",
        _df_to_markdown(data["hms_comparison_summary"]),
        "",
        _df_to_markdown(data["hms_comparison_metrics"]),
        "",
        data["hms_comparison_report_text"] or "unavailable",
        "",
        "## Charts",
        "",
    ]
    if data["charts"]:
        lines.extend([f"- `{_relative_or_absolute(path, data['project_dir'])}`" for path in data["charts"]])
    else:
        lines.append("unavailable")
    lines.extend(
        [
            "",
            "## Output File List",
            "",
            *[f"- `{_relative_or_absolute(path, data['project_dir'])}`" for path in data["assets"]],
            "",
            "## Missing or Unavailable Outputs",
            "",
            _df_to_markdown(data["missing_assets"]),
            "",
            "## Known Limits and Disclaimer",
            "",
            "- This report summarizes available local HydroLite Studio outputs and does not certify regulatory compliance.",
            "- Missing datasets or unavailable backends are shown as unavailable rather than inferred.",
            "- SWMM, GEE, and OpenHydroNet sections depend on optional local or cloud backends and credentials.",
        ]
    )
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return output


def _image_data_uri(path: Path) -> str:
    suffix = path.suffix.lower().lstrip(".") or "png"
    mime = "jpeg" if suffix in {"jpg", "jpeg"} else suffix
    return f"data:image/{mime};base64,{b64encode(path.read_bytes()).decode('ascii')}"


def render_project_report_html(project_dir: str | Path, output_path: str | Path | None = None) -> Path:
    data = collect_project_report_data(project_dir)
    output = Path(output_path).expanduser().resolve() if output_path else data["reports_dir"] / REPORT_FILENAMES["html"]
    output.parent.mkdir(parents=True, exist_ok=True)
    chart_html = ""
    if data["charts"]:
        chart_html = "\n".join(
            f"<figure><img src=\"{_image_data_uri(path)}\" alt=\"{escape(path.stem)}\"><figcaption>{escape(path.name)}</figcaption></figure>"
            for path in data["charts"]
        )
    else:
        chart_html = "<p class=\"muted\">unavailable</p>"
    sections = [
        ("Validation", _df_to_html(data["validation_overview"])),
        ("Scenario Runs", _df_to_html(data["batch_summary"])),
        ("Hydrology and Routing Comparison", _df_to_html(data["hydrology_metrics"])),
        ("Water Balance", _df_to_html(data["water_balance_metrics"])),
        ("SWMM Results", _df_to_html(data["swmm_metrics"])),
        ("HydroLite-SWMM Coupling", _df_to_html(data["coupling_metrics"])),
        ("GEE Data Center", _df_to_html(data["gee_summary"])),
        ("Observed Flow Evaluation", _df_to_html(data["model_performance"])),
        ("HEC-HMS and HydroLite Event Comparison", _df_to_html(data["hms_comparison_metrics"])),
        ("Missing or Unavailable Outputs", _df_to_html(data["missing_assets"])),
    ]
    html = [
        "<!doctype html>",
        "<html lang=\"en\"><head><meta charset=\"utf-8\"><title>HydroLite Project Report</title>",
        "<style>",
        "body{font-family:Arial,Helvetica,sans-serif;margin:40px;color:#1f2933;line-height:1.45} h1{color:#0f4c5c} h2{margin-top:30px;color:#1d3557} .meta,.muted{color:#667085} .summary{background:#f2f7f5;border-left:4px solid #2a9d8f;padding:12px 16px} table.data-table{border-collapse:collapse;width:100%;font-size:12px;margin:12px 0 20px} .data-table th{background:#eef2f7;text-align:left} .data-table th,.data-table td{border:1px solid #d0d5dd;padding:6px 8px;vertical-align:top} figure{margin:18px 0} img{max-width:100%;height:auto;border:1px solid #e5e7eb}",
        "</style></head><body>",
        f"<h1>{escape(data['project_name'])} Project Report</h1>",
        f"<p class=\"meta\">{escape(data['app_name'])} {escape(data['version'])} | Generated {escape(data['generated_at'])} | {escape(str(data['project_dir']))}</p>",
        "<h2>Executive Summary</h2><div class=\"summary\"><ul>",
        *[f"<li>{escape(line)}</li>" for line in _executive_summary(data)],
        "</ul></div>",
        "<h2>Project Overview</h2>",
        f"<p><strong>Project ID:</strong> {escape(data['project_id'])}<br><strong>Description:</strong> {escape(_stringify(data['project'].get('description', 'unavailable')))}</p>",
    ]
    for title, content in sections:
        html.extend([f"<h2>{escape(title)}</h2>", content])
    html.extend(["<h2>Charts</h2>", chart_html, "<h2>Known Limits and Disclaimer</h2>"])
    html.extend(
        [
            "<ul>",
            "<li>This report summarizes available local HydroLite Studio outputs and does not certify regulatory compliance.</li>",
            "<li>Missing datasets or unavailable backends are shown as unavailable rather than inferred.</li>",
            "<li>SWMM, GEE, and OpenHydroNet sections depend on optional local or cloud backends and credentials.</li>",
            "</ul></body></html>",
        ]
    )
    output.write_text("\n".join(html), encoding="utf-8")
    return output


def _docx_add_table(document: Any, df: pd.DataFrame, max_rows: int = 12) -> None:
    if df.empty:
        document.add_paragraph("unavailable")
        return
    small = df.head(max_rows).fillna("")
    table = document.add_table(rows=1, cols=len(small.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(small.columns):
        table.rows[0].cells[idx].text = str(column)
    for _, row in small.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(small.columns):
            cells[idx].text = _stringify(row[column])[:500]
    document.add_paragraph("")


def render_project_report_docx(project_dir: str | Path, output_path: str | Path | None = None) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:
        raise RuntimeError("python-docx is required to generate Word reports. Install package `python-docx`.") from exc

    data = collect_project_report_data(project_dir)
    output = Path(output_path).expanduser().resolve() if output_path else data["reports_dir"] / REPORT_FILENAMES["docx"]
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = f"{data['project_name']} Project Report"
    document.core_properties.author = data["app_name"]
    document.add_heading(f"{data['project_name']} Project Report", level=0)
    document.add_paragraph(f"{data['app_name']} {data['version']}")
    document.add_paragraph(f"Generated at: {data['generated_at']}")
    document.add_paragraph(f"Project path: {data['project_dir']}")
    document.add_heading("Executive Summary", level=1)
    for line in _executive_summary(data):
        document.add_paragraph(line, style="List Bullet")
    document.add_heading("Project Overview", level=1)
    document.add_paragraph(f"Project ID: {data['project_id']}")
    document.add_paragraph(f"Description: {_stringify(data['project'].get('description', 'unavailable'))}")
    if data["modules"]:
        document.add_heading("Enabled Modules", level=2)
        for name, enabled in data["modules"].items():
            document.add_paragraph(f"{name}: {enabled}", style="List Bullet")

    for title, key in [
        ("Validation", "validation_overview"),
        ("Scenario Runs", "batch_summary"),
        ("Hydrology and Routing Comparison", "hydrology_metrics"),
        ("Water Balance", "water_balance_metrics"),
        ("SWMM Results", "swmm_metrics"),
        ("HydroLite-SWMM Coupling", "coupling_metrics"),
        ("GEE Data Center", "gee_summary"),
        ("Observed Flow Evaluation", "model_performance"),
        ("HEC-HMS and HydroLite Event Comparison", "hms_comparison_metrics"),
        ("Missing or Unavailable Outputs", "missing_assets"),
    ]:
        document.add_heading(title, level=1)
        _docx_add_table(document, data[key])

    document.add_heading("Charts", level=1)
    if data["charts"]:
        for chart in data["charts"]:
            document.add_paragraph(chart.name)
            try:
                document.add_picture(str(chart), width=Inches(6.2))
            except Exception as exc:
                document.add_paragraph(f"Chart unavailable in Word export: {exc}")
    else:
        document.add_paragraph("unavailable")
    document.add_heading("Output File List", level=1)
    for asset in data["assets"][:80]:
        document.add_paragraph(_relative_or_absolute(asset, data["project_dir"]), style="List Bullet")
    document.add_heading("Known Limits and Disclaimer", level=1)
    for item in [
        "This report summarizes available local HydroLite Studio outputs and does not certify regulatory compliance.",
        "Missing datasets or unavailable backends are shown as unavailable rather than inferred.",
        "SWMM, GEE, and OpenHydroNet sections depend on optional local or cloud backends and credentials.",
    ]:
        document.add_paragraph(item, style="List Bullet")
    document.save(output)
    return output


def render_project_report_pdf(project_dir: str | Path, output_path: str | Path | None = None) -> Path:
    data = collect_project_report_data(project_dir)
    output = Path(output_path).expanduser().resolve() if output_path else data["reports_dir"] / REPORT_FILENAMES["pdf"]
    output.parent.mkdir(parents=True, exist_ok=True)
    html_path = render_project_report_html(project_dir)
    try:
        from weasyprint import HTML

        HTML(filename=str(html_path)).write_pdf(str(output))
        return output
    except Exception as exc:
        unavailable = data["reports_dir"] / REPORT_FILENAMES["pdf_unavailable"]
        unavailable.write_text(
            "\n".join(
                [
                    "# Project Report PDF Unavailable",
                    "",
                    "HydroLite generated Markdown, Word, and HTML reports, but PDF rendering is unavailable in this environment.",
                    "",
                    f"Requested PDF path: `{output}`",
                    f"HTML report path: `{html_path}`",
                    f"Reason: `{exc}`",
                    "",
                    "Install a supported PDF backend such as WeasyPrint to enable direct PDF export.",
                ]
            )
            + "\n",
            encoding="utf-8",
        )
        if output.exists():
            output.unlink()
        return unavailable


def export_project_report_bundle(project_dir: str | Path) -> Path:
    project = _project_path(project_dir)
    reports = _report_dir(project)
    render_project_report_markdown(project)
    render_project_report_docx(project)
    render_project_report_html(project)
    render_project_report_pdf(project)
    bundle = reports / REPORT_FILENAMES["bundle"]
    assets = list_report_assets(project)
    with zipfile.ZipFile(bundle, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for asset in assets:
            if asset == bundle or not _safe_bundle_member(asset, project):
                continue
            archive.write(asset, asset.relative_to(project))
    return bundle


def render_project_report_all(project_dir: str | Path) -> dict[str, Path]:
    markdown = render_project_report_markdown(project_dir)
    docx = render_project_report_docx(project_dir)
    html = render_project_report_html(project_dir)
    pdf_or_note = render_project_report_pdf(project_dir)
    bundle = export_project_report_bundle(project_dir)
    return {
        "markdown": markdown,
        "docx": docx,
        "html": html,
        "pdf": pdf_or_note,
        "bundle": bundle,
    }
